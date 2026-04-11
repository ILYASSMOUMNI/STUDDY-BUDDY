# backend/course_parser.py
# Parse les PDFs et DOCX uploadés → chunks de texte prêts pour l'indexation

import os
import re
from pathlib import Path
from typing import List, Dict
import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
    """Extrait tout le texte d'un PDF avec pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
    except Exception as e:
        print(f"[PARSER] Erreur PDF {filepath}: {e}")
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    """Extrait tout le texte d'un DOCX."""
    text = ""
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Inclure les tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text += row_text + "\n"
    except Exception as e:
        print(f"[PARSER] Erreur DOCX {filepath}: {e}")
    return text.strip()


def extract_text(filepath: str) -> str:
    """Dispatch selon l'extension du fichier."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Format non supporté : {ext}")


def clean_text(text: str) -> str:
    """Nettoie le texte extrait."""
    # Supprimer les lignes vides multiples
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Supprimer les espaces en début/fin de ligne
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    # Supprimer caractères parasites
    text = re.sub(r'[^\x00-\x7F\u00C0-\u024F\u0600-\u06FF\n .,;:!?\'\"()\[\]{}/\\@#$%&*+=<>_-]', '', text)
    return text.strip()


def split_into_chunks(text: str, chunk_size: int = 800, overlap: int = 150) -> List[Dict]:
    """
    Découpe le texte en chunks avec overlap.
    Retourne une liste de dicts: {chunk_id, text, char_start, char_end}
    """
    chunks = []
    words = text.split()
    total_words = len(words)

    if total_words == 0:
        return chunks

    step = chunk_size - overlap
    i = 0
    chunk_id = 0

    while i < total_words:
        chunk_words = words[i:i + chunk_size]
        chunk_text = " ".join(chunk_words)

        if len(chunk_text.strip()) > 50:  # ignorer les chunks trop courts
            chunks.append({
                "chunk_id": chunk_id,
                "text": chunk_text,
                "word_start": i,
                "word_end": i + len(chunk_words)
            })
            chunk_id += 1

        i += step

    return chunks


def parse_course_file(filepath: str, chunk_size: int = 800, overlap: int = 150) -> Dict:
    """
    Pipeline complet : fichier → texte nettoyé → chunks.
    Retourne:
      {
        "raw_text": str,
        "chunks": [{"chunk_id": int, "text": str, ...}],
        "num_chunks": int,
        "title": str  (nom du fichier sans extension)
      }
    """
    title = Path(filepath).stem.replace("_", " ").replace("-", " ").title()

    raw_text = extract_text(filepath)
    if not raw_text:
        raise ValueError(f"Impossible d'extraire du texte de {filepath}")

    cleaned = clean_text(raw_text)
    chunks = split_into_chunks(cleaned, chunk_size=chunk_size, overlap=overlap)

    print(f"[PARSER] '{title}' → {len(chunks)} chunks extraits")

    return {
        "title": title,
        "raw_text": cleaned,
        "chunks": chunks,
        "num_chunks": len(chunks)
    }


if __name__ == "__main__":
    # Test rapide
    import sys
    if len(sys.argv) > 1:
        result = parse_course_file(sys.argv[1])
        print(f"Titre: {result['title']}")
        print(f"Chunks: {result['num_chunks']}")
        print(f"\nPremier chunk:\n{result['chunks'][0]['text'][:300]}")
